import re
from pathlib import Path

import numpy as np
import pandas as pd
import plotly.express as px
import streamlit as st
from docx import Document

st.set_page_config(
    page_title="Nghiên cứu & Dự báo giá dầu DO 0,05S-II",
    page_icon="⛽",
    layout="wide",
    initial_sidebar_state="expanded",
)

DOCX_PATH = Path(__file__).with_name("bao_cao_gia_dau_do.docx")

CONST = -10270
BETA_CPI = 67590
BETA_USD = 0.5234
BETA_BRENT = 188.8192
R2 = 0.921
ADJ_R2 = 0.917

st.markdown(
    """
<style>
    .block-container {padding-top: 1.5rem; padding-bottom: 2rem;}
    .hero {
        padding: 36px 38px; border-radius: 28px;
        background: linear-gradient(135deg, #0f172a 0%, #1e3a5f 45%, #f59e0b 160%);
        color: white; margin-bottom: 22px;
    }
    .hero h1 {font-size: 44px; line-height: 1.12; margin: 0; font-weight: 900;}
    .hero p {font-size: 18px; opacity: .92; margin-top: 14px;}
    .card {
        background: #ffffff; border: 1px solid #e8edf5; border-radius: 20px;
        padding: 22px; box-shadow: 0 8px 28px rgba(15, 23, 42, 0.06); margin-bottom: 16px;
    }
    .dark-card {background:#0f172a; color:white; border-radius:20px; padding:22px; margin-bottom:16px;}
    .pill {display:inline-block; padding:7px 12px; border-radius:999px; background:#eef6ff; color:#0f3b66; font-weight:700; margin: 2px 5px 8px 0;}
    .section-title {font-size: 30px; font-weight: 900; color:#0f172a; margin: 6px 0 16px 0;}
    .muted {color:#64748b;}
    .chapter-box h3 {margin-top: 20px;}
    .chapter-box p {text-align: justify; font-size: 16px; line-height: 1.75;}
    .stMetric {background:#fff; border:1px solid #e8edf5; border-radius:16px; padding:12px;}
</style>
""",
    unsafe_allow_html=True,
)

@st.cache_data(show_spinner=False)
def load_docx(path: str):
    doc = Document(path)
    items = []
    for p in doc.paragraphs:
        text = " ".join(p.text.split())
        if text:
            style = p.style.name if p.style else ""
            items.append({"type": "p", "text": text, "style": style})
    for table in doc.tables:
        rows = []
        for row in table.rows:
            rows.append([" ".join(cell.text.split()) for cell in row.cells])
        if rows:
            items.append({"type": "table", "rows": rows})
    return items

items = load_docx(str(DOCX_PATH)) if DOCX_PATH.exists() else []
texts = [x["text"] for x in items if x["type"] == "p"]
full_text = "\n".join(texts)

TEAM = pd.DataFrame({
    "STT": [1, 2, 3, 4, 5],
    "Họ tên": ["Lê Thái Vỹ", "Lưu Huỳnh Nhất", "Lê Ngọc Thịnh", "Nguyễn Huy Hiệu", "Ngô Sử Học"],
    "Chức vụ": ["Trưởng nhóm", "Thành viên", "Thành viên", "Thành viên", "Thành viên"],
    "Mức độ hoàn thành": ["10/10", "10/10", "10/10", "10/10", "10/10"],
})

PAGES = [
    "🏠 Trang chủ",
    "👥 Thông tin nhóm",
    "📌 Tóm tắt báo cáo",
    "🚀 Phần mở đầu",
    "📚 Chương 1",
    "🧠 Chương 2",
    "📈 Chương 3",
    "✅ Kết luận",
    "📎 Phụ lục & tài liệu",
    "⛽ Dự báo OLS",
]


def clean_heading(s: str) -> str:
    return re.sub(r"\s+", " ", s.strip())


def section_between(start_patterns, stop_patterns=None, max_chars=30000):
    start_idx = None
    for i, t in enumerate(texts):
        if any(re.search(p, t) for p in start_patterns):
            start_idx = i
            break
    if start_idx is None:
        return []
    end_idx = len(texts)
    if stop_patterns:
        for j in range(start_idx + 1, len(texts)):
            if any(re.search(p, texts[j]) for p in stop_patterns):
                end_idx = j
                break
    out = []
    total = 0
    for t in texts[start_idx:end_idx]:
        total += len(t)
        if total > max_chars:
            out.append("[Nội dung còn lại xem trong file Word đính kèm khi triển khai.]")
            break
        out.append(t)
    return out


def render_text_block(lines, limit=None):
    if not lines:
        st.info("Chưa đọc được nội dung phần này từ file Word.")
        return
    if limit:
        lines = lines[:limit]
    st.markdown('<div class="chapter-box">', unsafe_allow_html=True)
    for line in lines:
        if re.match(r"^(CHƯƠNG|PHẦN MỞ ĐẦU|KẾT LUẬN|TÓM TẮT|TÀI LIỆU|PHỤ LỤC)", line, re.I):
            st.markdown(f"### {line}")
        elif re.match(r"^\d+(\.\d+)*\.?\s", line):
            st.markdown(f"#### {line}")
        elif len(line) < 90 and (line.isupper() or line.startswith(("LỜI", "DANH MỤC"))):
            st.markdown(f"### {line}")
        else:
            st.write(line)
    st.markdown('</div>', unsafe_allow_html=True)


def top_metrics():
    a, b, c, d = st.columns(4)
    a.metric("Số quan sát", "72 tháng")
    b.metric("R²", f"{R2:.3f}")
    c.metric("R² hiệu chỉnh", f"{ADJ_R2:.3f}")
    d.metric("Biến độc lập", "CPI, USD/VND, Brent")


def page_home():
    st.markdown(
        """
<div class="hero">
    <h1>Các yếu tố ảnh hưởng đến giá dầu DO 0,05S-II tại Việt Nam</h1>
    <p>Website trình bày tiểu luận Smart Logistics theo dạng báo cáo số, kết hợp mô hình hồi quy tuyến tính đa biến OLS và công cụ dự báo tương tác.</p>
</div>
""",
        unsafe_allow_html=True,
    )
    top_metrics()
    c1, c2 = st.columns([1.2, 1])
    with c1:
        st.markdown('<div class="card">', unsafe_allow_html=True)
        st.subheader("Thông tin đề tài")
        st.write("**Trường:** Đại học Giao thông Vận tải TP. Hồ Chí Minh")
        st.write("**Viện:** Viện Hàng hải")
        st.write("**Môn:** Smart Logistics")
        st.write("**Giảng viên hướng dẫn:** TS. Trịnh Tuấn Hùng")
        st.write("**Đơn vị thực hiện:** Nhóm 14")
        st.write("**Thời gian:** TP. Hồ Chí Minh – 07/2026")
        st.markdown('</div>', unsafe_allow_html=True)
    with c2:
        st.markdown('<div class="dark-card">', unsafe_allow_html=True)
        st.subheader("Cấu trúc website")
        st.write("Trang chủ → Thông tin nhóm → Tóm tắt → Mở đầu → Chương 1 → Chương 2 → Chương 3 → Kết luận → Phụ lục → Dự báo OLS.")
        st.markdown('</div>', unsafe_allow_html=True)


def page_team():
    st.markdown('<div class="section-title">👥 Thông tin nhóm</div>', unsafe_allow_html=True)
    st.dataframe(TEAM, use_container_width=True, hide_index=True)
    st.markdown('<div class="card">', unsafe_allow_html=True)
    st.subheader("Bảng phân công")
    st.write("Có thể bổ sung nội dung thực hiện chi tiết cho từng thành viên ngay tại bảng này trước khi nộp hoặc deploy website.")
    st.markdown('</div>', unsafe_allow_html=True)


def page_summary():
    st.markdown('<div class="section-title">📌 Tóm tắt báo cáo</div>', unsafe_allow_html=True)
    render_text_block(section_between([r"^TÓM TẮT BÁO CÁO"], [r"^LỜI CAM ĐOAN"]), limit=40)


def page_intro():
    st.markdown('<div class="section-title">🚀 Phần mở đầu</div>', unsafe_allow_html=True)
    render_text_block(section_between([r"^PHẦN MỞ ĐẦU"], [r"^CHƯƠNG 1"]), limit=120)


def page_chapter1():
    st.markdown('<div class="section-title">📚 Chương 1</div>', unsafe_allow_html=True)
    render_text_block(section_between([r"^CHƯƠNG 1"], [r"^CHƯƠNG 2"]), limit=180)


def page_chapter2():
    st.markdown('<div class="section-title">🧠 Chương 2</div>', unsafe_allow_html=True)
    st.markdown("""
<div class="card">
<span class="pill">CPI</span><span class="pill">USD/VND</span><span class="pill">Brent</span><span class="pill">OLS</span>
<br><br><b>Mô hình nghiên cứu:</b> Giá DO = β0 + β1*CPI + β2*USD/VND + β3*Brent + ε
</div>
""", unsafe_allow_html=True)
    render_text_block(section_between([r"^CHƯƠNG 2"], [r"^CHƯƠNG 3"]), limit=220)


def page_chapter3():
    st.markdown('<div class="section-title">📈 Chương 3</div>', unsafe_allow_html=True)
    top_metrics()
    coef = pd.DataFrame({
        "Biến": ["Hằng số", "CPI", "USD/VND", "Brent"],
        "Hệ số": [CONST, BETA_CPI, BETA_USD, BETA_BRENT],
        "Diễn giải nhanh": [
            "Mức chặn của mô hình",
            "CPI tăng 1 đơn vị làm giá DO tăng theo hệ số này",
            "Tỷ giá tăng 1 VND/USD làm giá DO tăng 0,5234 đồng/lít",
            "Brent tăng 1 USD/thùng làm giá DO tăng khoảng 188,82 đồng/lít",
        ],
    })
    st.dataframe(coef, use_container_width=True, hide_index=True)
    render_text_block(section_between([r"^CHƯƠNG 3"], [r"^KẾT LUẬN CHUNG", r"^KẾT LUẬN"]), limit=260)


def page_conclusion():
    st.markdown('<div class="section-title">✅ Kết luận</div>', unsafe_allow_html=True)
    render_text_block(section_between([r"^KẾT LUẬN"], [r"^TÀI LIỆU THAM KHẢO", r"^PHỤ LỤC"]), limit=120)


def page_appendix():
    st.markdown('<div class="section-title">📎 Phụ lục & tài liệu</div>', unsafe_allow_html=True)
    render_text_block(section_between([r"^TÀI LIỆU THAM KHẢO", r"^PHỤ LỤC"], None, max_chars=45000), limit=350)


def page_predictor():
    st.markdown('<div class="section-title">⛽ Hệ thống dự báo giá dầu DO 0,05S-II</div>', unsafe_allow_html=True)
    st.caption("Mô hình hồi quy tuyến tính đa biến OLS dựa trên CPI, tỷ giá USD/VND và giá dầu Brent.")
    st.code("Giá DO = -10270 + 67590*CPI + 0.5234*USD/VND + 188.8192*Brent")
    col1, col2, col3 = st.columns(3)
    with col1:
        cpi = st.number_input("CPI", min_value=-0.05, max_value=0.20, value=0.035, step=0.001, format="%.4f")
    with col2:
        usd_vnd = st.number_input("Tỷ giá USD/VND", min_value=15000, max_value=35000, value=24500, step=100)
    with col3:
        brent = st.number_input("Giá dầu Brent (USD/thùng)", min_value=0.0, max_value=200.0, value=80.0, step=1.0)

    gia_du_bao = CONST + BETA_CPI * cpi + BETA_USD * usd_vnd + BETA_BRENT * brent
    k1, k2, k3, k4 = st.columns(4)
    k1.metric("Giá DO dự báo", f"{gia_du_bao:,.0f} đồng/lít")
    k2.metric("CPI đầu vào", f"{cpi*100:.2f}%")
    k3.metric("USD/VND đầu vào", f"{usd_vnd:,.0f}")
    k4.metric("Brent đầu vào", f"{brent:.2f} USD/thùng")

    if gia_du_bao < 18000:
        st.success("Giá dự báo ở vùng thấp.")
    elif gia_du_bao <= 23000:
        st.info("Giá dự báo ở vùng trung bình.")
    else:
        st.warning("Giá dự báo ở vùng cao, doanh nghiệp nên chú ý rủi ro chi phí nhiên liệu.")

    contrib = pd.DataFrame({
        "Yếu tố": ["Hằng số", "CPI", "USD/VND", "Brent"],
        "Giá trị đóng góp": [CONST, BETA_CPI * cpi, BETA_USD * usd_vnd, BETA_BRENT * brent],
    })
    fig_contrib = px.bar(contrib, x="Yếu tố", y="Giá trị đóng góp", text="Giá trị đóng góp", title="Đóng góp của từng thành phần vào giá dự báo")
    fig_contrib.update_traces(texttemplate="%{text:,.0f}", textposition="outside")
    st.plotly_chart(fig_contrib, use_container_width=True)

    brent_range = np.arange(50, 121, 5)
    scenario = pd.DataFrame({"Brent": brent_range})
    scenario["Giá DO dự báo"] = CONST + BETA_CPI * cpi + BETA_USD * usd_vnd + BETA_BRENT * scenario["Brent"]
    fig_sce = px.line(scenario, x="Brent", y="Giá DO dự báo", markers=True, title="Giá DO thay đổi khi Brent biến động")
    st.plotly_chart(fig_sce, use_container_width=True)

    with st.expander("Xem bảng mô phỏng kịch bản"):
        st.dataframe(scenario, use_container_width=True)

    st.subheader("📂 Tải dữ liệu CSV để vẽ biểu đồ thực tế")
    st.write("File CSV nên có các cột: `Thang`, `CPI`, `USD_VND`, `Gia_Brent`, `Gia_DO_005S_II`")
    uploaded_file = st.file_uploader("Tải file CSV dữ liệu 2020–2025", type=["csv"])
    if uploaded_file is not None:
        df = pd.read_csv(uploaded_file)
        required_cols = ["Thang", "CPI", "USD_VND", "Gia_Brent", "Gia_DO_005S_II"]
        missing_cols = [c for c in required_cols if c not in df.columns]
        if missing_cols:
            st.error(f"File đang thiếu cột: {missing_cols}")
        else:
            df["Gia_Du_Bao"] = CONST + BETA_CPI * df["CPI"] + BETA_USD * df["USD_VND"] + BETA_BRENT * df["Gia_Brent"]
            st.success("Đã đọc dữ liệu thành công.")
            fig_real = px.line(df, x="Thang", y=["Gia_DO_005S_II", "Gia_Du_Bao"], title="So sánh giá dầu thực tế và giá dự báo", markers=True)
            st.plotly_chart(fig_real, use_container_width=True)
            fig_vars = px.line(df, x="Thang", y=["CPI", "Gia_Brent"], title="Diễn biến CPI và giá Brent")
            st.plotly_chart(fig_vars, use_container_width=True)
            st.dataframe(df, use_container_width=True)
            csv = df.to_csv(index=False).encode("utf-8-sig")
            st.download_button("⬇️ Tải dữ liệu có cột dự báo", data=csv, file_name="du_lieu_du_bao_gia_dau.csv", mime="text/csv")


with st.sidebar:
    st.title("⛽ Oil DO Research")
    page = st.radio("Chọn trang", PAGES, label_visibility="collapsed")
    st.divider()
    st.caption("Nhóm 14 · Smart Logistics · OLS")

if page == "🏠 Trang chủ":
    page_home()
elif page == "👥 Thông tin nhóm":
    page_team()
elif page == "📌 Tóm tắt báo cáo":
    page_summary()
elif page == "🚀 Phần mở đầu":
    page_intro()
elif page == "📚 Chương 1":
    page_chapter1()
elif page == "🧠 Chương 2":
    page_chapter2()
elif page == "📈 Chương 3":
    page_chapter3()
elif page == "✅ Kết luận":
    page_conclusion()
elif page == "📎 Phụ lục & tài liệu":
    page_appendix()
elif page == "⛽ Dự báo OLS":
    page_predictor()
